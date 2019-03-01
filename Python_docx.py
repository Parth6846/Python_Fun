# -*- coding: utf-8 -*-
"""
Created on Wed Feb 27 15:45:43 2019

@author: parthra5
"""

from openpyxl import Workbook
from docx import Document
import os
import openpyxl as op
import re

def main():
    dir_n="C:\\Users\\PARTHRA5\\OneDrive - Novartis Pharma AG\\work 28-2\\"
    docs="BPS"
    files=os.listdir(dir_n+docs)
    #files=['ISRM_IAM_APS_TS_001']
    wb_n=dir_n+"B.xlsx"
    
    xls=op.load_workbook(wb_n)
    ws1=xls['Sheet1']
    j=1
    head=['Subject','test name	','Description',	'Test Case Type','Step Name (Design Steps)',	'Description (Design Steps)',	'Expected Result (Design Steps)',	'Verification Point',	'Approval Status','Designer','Type',	'Reviewer1	','Reviewer2',	'Document Id',	'Document Name',	'Responsible Owner']
    for f in files:
        try:
            ws1.cell(row=j,column=1).value=f
            docid=f[:17]
            ws=xls.create_sheet(docid)
            doc=Document(dir_n+docs+'\\'+f)
            tables=doc.tables
            desc=[]
            for i in range(1,len(head)):
                    ws.cell(row=1,column=i).value=head[i-1]
            
            
            t=0
            
            while(t<len(tables)):
                if("TEST EXECUTION" in doc.tables[t].rows[0].cells[0].text):
                    dtable=tables[t]
                    row_count=len(dtable.rows)
                    
                    for r in range(2,row_count):                    
                        s=""                                        
                        for c in range(1,4):                         
                             #print("parth:"+s)
                             #print("sa:"+dtable.rows[r].cells[c].text)
                             if(c==2):
                                continue
                             s=s+dtable.rows[r].cells[c].text+"\n"                     
                        desc.append(s)
                    print(desc)
                                                
                t=t+1
            
                  
            
            t=0
            r_n=2
            
            Table_count=0
            
            
            
            while(t<len(tables)):
                if("TC/ Step" in doc.tables[t].rows[0].cells[0].text and "Test Description / Instruction" in doc.tables[t].rows[0].cells[1].text):
                   Table_count=Table_count+1
                t=t+1
            
            
            t=0
            if(Table_count==len(desc)):
               print("success")
               cn=0
               cp=1
               while(t<len(tables)):
                
                   if("TC/ Step" in doc.tables[t].rows[0].cells[0].text and "Test Description / Instruction" in doc.tables[t].rows[0].cells[1].text):
                        print("success 1")
                        dtable=tables[t]
                        ts=0
                                               
                        
                        descp=desc[cn]
                        inx=descp.find('.')
                       
                        
                        ja=descp[0:inx]
                        
                        ja=re.sub(r'Test Case Name & Objective:','',ja)
                        ja=re.sub(r'\n','',ja)
                        
                        
                        for r in dtable.rows:
                            if(ts==0):
                                ts=ts+1
                                continue                         
                            
                            ws.cell(row=r_n,column=2).value=docid+"_TC00"+str(cp)+"_"+ja
                            ws.cell(row=r_n,column=3).value=descp
                            ws.cell(row=r_n,column=4).value="OQ"
                            ws.cell(row=r_n,column=5).value=r.cells[0].text
                            ws.cell(row=r_n,column=6).value=r.cells[1].text
                            ws.cell(row=r_n,column=7).value=r.cells[2].text
                            ws.cell(row=r_n,column=8).value="Y"
                            ws.cell(row=r_n,column=9).value="In Progress"
                            ws.cell(row=r_n,column=11).value="MANUAL"
                            ws.cell(row=r_n,column=14).value=docid
                            r_n=r_n+1
                            ts=ts+1                            
                            flag=True
                        cn=cn+1
                        cp=cp+1
                    
                   t=t+1       
               if(flag==True):
                    ws1.cell(row=j,column=2).value="done"
               else:
                    ws1.cell(row=j,column=2).value="different format"
                    xls.remove_sheet(ws)
               j=j+1      
            else:
                ws1.cell(row=j,column=2).value="Table Count is different"
                xls.remove_sheet(ws)       
            #j=2
            #for i in range(0, len(desc)):
             #   ws.cell(row=j,column=1).value=desc[i]
              #  j=j+1
               # i=i+1
            
                     
                   
                
            #print(desc)
            
            
                
                    
            
            
            j=j+1
            
        except Exception as e: 
                               print("Exception")
                               print(e)
                               ws1.cell(row=j,column=2).value="some error"
                               j=j+1
    xls.save(wb_n)
if __name__== "__main__":
    main()
